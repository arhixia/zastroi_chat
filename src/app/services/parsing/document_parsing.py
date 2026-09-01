"""
Извлечение текста из документов, загруженных администратором в базу знаний
Каждая функция возвращает "сырой" текст документа —
дальше он идёт в indexing.index_text() точно так же, как текст страницы сайта.

Если из документа не удалось извлечь текст (скан без текстового слоя,
пустой файл) — extract_text() просто вернёт пустую строку. Документ при
этом всё равно сохраняется в БД как обычно: index_text() на пустом тексте
создаст 0 чанков, и бот такой источник просто не найдёт при retrieval —
никакой отдельной обработки ошибки не требуется, алгоритм и так его не заденет.
"""
from pathlib import Path

import openpyxl
import pdfplumber
from docx import Document as DocxDocument

from app.db.models import DocumentType


def extract_text_from_pdf(path: Path) -> str:
    """
    Извлекает текст из PDF постранично. Если у PDF нет текстового слоя
    (скан — просто картинка), pdfplumber вернёт пустые строки на каждой
    странице, итоговый текст будет пустым.
    """
    pages_text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            pages_text.append(page.extract_text() or "")
    return "\n\n".join(pages_text)


def extract_text_from_docx(path: Path) -> str:
    """
    Извлекает текст из .docx: абзацы + текст из таблиц (прайсы часто
    оформлены таблицей, а не абзацами).
    """
    doc = DocxDocument(path)

    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def extract_text_from_xlsx(path: Path) -> str:
    """
    Превращает таблицу (прайс-лист, характеристики квартир) в читаемый текст.
    Первая строка листа считается заголовками — каждая следующая строка
    превращается в "заголовок: значение; заголовок: значение", чтобы модель
    видела не голые числа, а что они означают (площадь: 45; цена: 6500000).
    """
    workbook = openpyxl.load_workbook(path, data_only=True)
    parts: list[str] = []

    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue

        parts.append(f"Лист: {sheet.title}")
        headers = [str(h).strip() if h is not None else "" for h in rows[0]]
        has_headers = any(headers)

        for row in rows[1:] if has_headers else rows:
            values = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if not values:
                continue

            if has_headers:
                pairs = [
                    f"{headers[i]}: {v}"
                    for i, v in enumerate(row)
                    if v is not None and str(v).strip() and i < len(headers) and headers[i]
                ]
                parts.append("; ".join(pairs) if pairs else " | ".join(values))
            else:
                parts.append(" | ".join(values))

    return "\n".join(parts)


_EXTRACTORS = {
    DocumentType.pdf: extract_text_from_pdf,
    DocumentType.docx: extract_text_from_docx,
    DocumentType.txt: extract_text_from_txt,
    DocumentType.xlsx: extract_text_from_xlsx,
}


def extract_text(path: Path, file_type: DocumentType) -> str:
    """
    Диспетчер: вызывает нужный экстрактор по типу файла.
    Может вернуть пустую строку — это штатный исход для скана без текста,
    вызывающий код (загрузка документа) ничего специального с этим не делает.
    """
    extractor = _EXTRACTORS.get(file_type)
    if extractor is None:
        raise NotImplementedError(f"Извлечение текста для {file_type} файла пока не реализовано")

    return extractor(path)